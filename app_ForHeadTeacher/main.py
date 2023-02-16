# -*- coding: utf-8 -*-
import sqlite3 as sq
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def sq_init():
    with sq.connect("school.db") as db:
        cur = db.cursor()
        #Создание пустой базы данных
        query = """
        CREATE TABLE IF NOT EXISTS n_lesson(
            n_lesson INTEGER 
        );
        CREATE TABLE IF NOT EXISTS day(
            day TEXT 
        );
        CREATE TABLE IF NOT EXISTS grade(
            grade TEXT
        );
        CREATE TABLE IF NOT EXISTS id_subjects(
            id_subject INTEGER NOT NULL,
            subject TEXT,
            PRIMARY KEY("id_subject" AUTOINCREMENT)
        );
        CREATE TABLE IF NOT EXISTS id_students(
            id_student INTEGER NOT NULL,
            name_1 TEXT,
            name_2 TEXT,
            grade TEXT,
            PRIMARY KEY("id_student" AUTOINCREMENT)
        );
        CREATE TABLE IF NOT EXISTS id_teachers(
            id_teacher INTEGER NOT NULL,
            name_1 TEXT,
            name_2 TEXT,
            name_3 TEXT,
            fix_cl_room TEXT,
            PRIMARY KEY("id_teacher" AUTOINCREMENT)
        );
        CREATE TABLE IF NOT EXISTS students(
            id_student INTEGER,
            subject TEXT,
            mark INTEGER
        );
        CREATE TABLE IF NOT EXISTS teachers(
            id_teacher INTEGER,
            subject TEXT,
            grade TEXT,
            day TEXT,
            n_lesson INTEGER,
            room TEXT
        )
        """
        cur.executescript(query)
def sq_dump_w():
    with sq.connect("school.db") as db:
        cur = db.cursor()
        with open("sql_dump.sql", "w") as f:
            for sql in db.iterdump():
                f.write(sql)
def sq_dump_r():
    with sq.connect("school.db") as db:
        cur = db.cursor()
        with open("sql_dump.sql", "r") as f:
            sql = f.read()
            cur.executescript(sql)

def choose_user():
    os.system('cls')
    print("Укажите ваш класс пользователя:\n")
    print("1 - Учитель")
    print("2 - Заведующий учебной работой")
def choose_main(user):
    os.system('cls')
    print("Введите .. чтобы выбрать действие:\n")
    print("1 - Исправить оценку ученику")
    if user == '2':
        print("2 - Редактировать информацию об ученике/учителе")
    print("3 - Какие предметы будут в ?? классе в ?? день")
    if user == '2':
        print("4 - Кто из учителей преподает в ?? классе")
    print("5 - В каком кабинете будет ??-й урок в ?? у ?? класса")
    if user == '2':
        print("6 - В каких классах преподает ?? предмет ??")
        print("7 - Формирование справки о количестве учеников")
        print("8 - Формирование отчета о работе школы")
        print("9 - Создание резервной копии БД")
    print("0 - Выход\n")
def choose_key2_1():
    os.system('cls')
    print("Чью информацию будем редактировать?\n")
    print("1 - Ученик") #49
    print("2 - Учитель") #50
    print("0 - Вернуться назад\n") #48
def choose_key2_2(key):
    os.system('cls')
    if key == 1:
        print("Что будем делать с информацией об ученике?\n")
    if key == 2:
        print("Что будем делать с информацией об учителе?\n")
    print("1 - Добавление") #49
    print("2 - Удаление") #50
    print("0 - Вернуться назад\n") #48

def key1():
    try:
        db = sq.connect("school.db")
        cur = db.cursor()
        cur.execute("""Begin;""")
        print("||  Введите фамилию и имя ученика, оценку которого нужно исправить")
        name_1 = input("  ")
        name_2 = input("  ")
        request = """
            SELECT name_2, subject, mark
            FROM students
            JOIN id_students ON id_students.id_student = students.id_student
            WHERE students.id_student =
            (SELECT id_student FROM id_students
            WHERE name_1 = ? AND name_2 = ?);
        """
        cur.execute(request, (name_1, name_2))
        output = cur.fetchall()
        #print (output)
        if output == []:
            print("||  Такого ученика нет")
        else:
            for row in output:
                print ("\t", '%-20s %-15d' % (row[1], row[2])) 
            print("\n||  Какой предмет исправляем?")
            subject = input("  ")        
            request = """
                SELECT name_2 as 'ученик', subject as 'предмет', mark as 'оценка'
                FROM students
                JOIN id_students ON id_students.id_student = students.id_student
                WHERE students.id_student =
                (SELECT id_student FROM id_students
                WHERE name_1 = ? AND name_2 = ?)
                AND subject = ?;
            """
            cur.execute(request, (name_1, name_2, subject))
            output = cur.fetchall()
            #print (output)
            if output == []:
                print("||  Некоректный ввод или оценок нет, исправлять нечего")
            else:
                for row in output:
                    print ("\t", '%-20s' % row[2]) 
                print("\n||  Какую оценку исправляем и на какую?")
                mark_to = input("  ")
                mark = input("  ")
                request = """
                    UPDATE students SET mark = ?
                    WHERE (students.id_student = (SELECT id_student FROM id_students
                    WHERE name_1 = ? AND name_2 = ?))
                    AND (rowid = (SELECT ROWID FROM students
                    WHERE mark = ? AND (students.id_student = (SELECT id_student FROM id_students
                    WHERE name_1 = ? AND name_2 = ?)) AND (subject = ?)));
                """
                cur.execute(request, (mark, name_1, name_2, mark_to, name_1, name_2, subject))
                db.commit()
                print("||  Оценка исправлена!")
        os.system('pause')

    except sq.Error as e:
        if db: db.rollback()
        print("||  Ошибка выполнения запроса ", e)
        os.system('pause')
    finally:
        if db: db.close()
def key2():
    choose_key2_1()
    key = input()
    while key != '0':
        choose_key2_1()
        if key == '1':
            os.system('cls')
            while key != '0':#Ученик
                choose_key2_2(1)
                key = input()
                if key == '1':#Добавление
                    os.system('cls')
                    try:
                        db = sq.connect("school.db")
                        cur = db.cursor()
                        cur.execute("""Begin;""")
                        print("||  Введите класс в который добавляем ученика")
                        grade = input("  ")
                        request = """
                            SELECT grade FROM grade WHERE grade = ?;
                        """
                        cur.execute(request, (grade,))
                        output = cur.fetchall()
                        #print (output)
                        if output == []:
                            print("||  Такого класса в школе нет")
                        else:
                            print("||  Введите фамилию и имя ученика которого добавляем в", grade)
                            name_1 = input("  ")  
                            name_2 = input("  ")
                            request = "INSERT INTO id_students (name_1, name_2, grade) VALUES (?, ?, ?);"
                            cur.execute(request, (name_1, name_2, grade,))
                            print("||  Ученик добавлен!")
                            request = "SELECT * FROM id_students WHERE name_1 = ? AND name_2 = ? AND grade = ?;"
                            cur.execute(request, (name_1, name_2, grade,))
                            output = cur.fetchall()
                            #print (output)
                            for row in output:
                                print ("\t", '%-20s %-15s %-15s' % (row[1],row[2],row[3]))
                            db.commit()
                        os.system('pause')
                    except sq.Error as e:
                        if db: db.rollback()
                        print("||  Ошибка выполнения запроса ", e)
                        os.system('pause')
                    finally:
                        if db: db.close()
                if key == '2':#Удаление
                    os.system('cls')          
                    try:
                        db = sq.connect("school.db")
                        cur = db.cursor()
                        cur.execute("""Begin;""")
                        print("||  Введите класс из которого удаляем ученика")
                        grade = input("  ")
                        request = """
                            SELECT grade FROM grade WHERE grade = ?;
                        """
                        cur.execute(request, (grade,))
                        output = cur.fetchall()
                        #print (output)
                        if output == []:
                            print("||  Такого класса в школе нет")
                        else:
                            print("||  Введите фамилию и имя ученика в", grade, "которого удаляем")
                            name_1 = input("  ")  
                            name_2 = input("  ")
                            request = """
                                DELETE FROM id_students WHERE name_1 = ? AND name_2 = ? AND grade = ?
                                AND rowid = (SELECT ROWID FROM id_students WHERE name_1 = ? AND name_2 = ? AND grade = ?);
                            """
                            cur.execute(request, (name_1, name_2, grade, name_1, name_2, grade,))
                            print("||  Ученик удален")
                            db.commit()
                        os.system('pause')
                    except sq.Error as e:
                        if db: db.rollback()
                        print("||  Ошибка выполнения запроса ", e)
                        os.system('pause')
                    finally:
                        if db: db.close()
            choose_key2_1()
            key = input()
        if key == '2':
            os.system('cls')
            while key != '0':#Учитель
                choose_key2_2(2)
                key = input()
                if key == '1':#Добавление
                    os.system('cls')
                    try:
                        db = sq.connect("school.db")
                        cur = db.cursor()
                        cur.execute("""Begin;""")
                        print("||  Введите ФИО нового учителя")
                        name_1 = input("  ")
                        name_2 = input("  ")
                        name_3 = input("  ")
                        print("||  Закреплен ли за ним кабинет?\n||  0 - Нет\n||  1 - Да")
                        _if = input("  ")
                        if _if == '1':
                            print("||  Введите номер кабинета")
                            fix_cl_room = input()
                            request = "INSERT INTO id_teachers (name_1, name_2, name_3, fix_cl_room) VALUES (?, ?, ?, ?);"
                            cur.execute(request, (name_1, name_2, name_3, fix_cl_room,))
                            print("||  Учитель добавлен!")
                            request = "SELECT * FROM id_teachers WHERE name_1 = ? AND name_2 = ? AND name_3 = ? AND fix_cl_room = ?;"
                            cur.execute(request, (name_1, name_2, name_3, fix_cl_room,))
                            output = cur.fetchall()
                            #print (output)
                            for row in output:
                                print ("\t", '%-15s %-13s %-18s %-10s' % (row[1],row[2],row[3],row[4]))
                        if _if == '0':
                            request = "INSERT INTO id_teachers (name_1, name_2, name_3) VALUES (?, ?, ?);"
                            cur.execute(request, (name_1, name_2, name_3,))
                            print("||  Учитель добавлен!")
                            request = "SELECT * FROM id_teachers WHERE name_1 = ? AND name_2 = ? AND name_3 = ?;"
                            cur.execute(request, (name_1, name_2, name_3,))
                            output = cur.fetchall()
                            #print (output)
                            for row in output:
                                print ("\t", '%-15s %-13s %-18s' % (row[1],row[2],row[3]))
                        db.commit()
                        os.system('pause')
                    except sq.Error as e:
                        if db: db.rollback()
                        print("||  Ошибка выполнения запроса ", e)
                        os.system('pause')
                    finally:
                        if db: db.close()
                    choose_key2_2(2)
                    key = input()
                if key == '2':#Удаление
                    os.system('cls')          
                    try:
                        db = sq.connect("school.db")
                        cur = db.cursor()
                        cur.execute("""Begin;""")
                        print("||  Введите ФИО учителя")
                        name_1 = input("  ")
                        name_2 = input("  ")
                        name_3 = input("  ")
                        request = """
                            DELETE FROM id_teachers WHERE name_1 = ? AND name_2 = ? AND name_3 = ?
                            AND rowid = (SELECT ROWID FROM id_teachers WHERE name_1 = ? AND name_2 = ? AND name_3 = ?);
                        """
                        cur.execute(request, (name_1, name_2, name_3, name_1, name_2, name_3,))
                        print("||  Учитель удален")
                        db.commit()
                        os.system('pause')
                    except sq.Error as e:
                        if db: db.rollback()
                        print("||  Ошибка выполнения запроса ", e)
                        os.system('pause')
                    finally:
                        if db: db.close()
                    choose_key2_2(2)
                    key = input()
            choose_key2_1()
            key = input()
    return 0
def key3():
    try:
        db = sq.connect("school.db")
        cur = db.cursor()
        cur.execute("""Begin;""")        
        print("||  В какой день(ДД.ММ.ГГГГ) будет урок?")
        day = input("  ")
        print("||  У какого класса?")
        grade = input("  ")
        request = """
            SELECT n_lesson, subject, room, name_1, name_2, name_3
            FROM teachers
            JOIN id_teachers ON id_teachers.id_teacher = teachers.id_teacher
            WHERE grade = ? AND day = ?
        """
        cur.execute(request, (grade, day,))
        #print (output)
        output = cur.fetchall()
        if output == []:
            print("||  Такого расписания нет")
        else:
            for row in output:
                print ("\t", '%-3s %-20s %-4s %-18s %-13s %-20s' % (row[0],row[1],row[2],row[3],row[4],row[5])) 
        os.system('pause')
    except sq.Error as e:
        if db: db.rollback()
        print("||  Ошибка выполнения запроса ", e)
        os.system('pause')
    finally:
        if db: db.close()
def key4():
    try:
        db = sq.connect("school.db")
        cur = db.cursor()
        cur.execute("""Begin;""")        
        print("||  Введите класс для которого нужно вывести список преподавателей")
        grade = input("  ")
        request = """
            SELECT DISTINCT subject, name_1, name_2, name_3
            FROM teachers
            JOIN id_teachers ON id_teachers.id_teacher = teachers.id_teacher
            WHERE grade = ?
        """
        cur.execute(request, (grade,))
        #print (output)
        output = cur.fetchall()
        if output == []:
            print("||  Преподаватели для этого класса еще не назначены")
        else:
            for row in output:
                print ("\t", '%-20s %-18s %-13s %-20s' % (row[0],row[1],row[2],row[3])) 
        os.system('pause')
    except sq.Error as e:
        if db: db.rollback()
        print("||  Ошибка выполнения запроса ", e)
        os.system('pause')
    finally:
        if db: db.close()
def key5():
    try:
        db = sq.connect("school.db")
        cur = db.cursor()
        cur.execute("""Begin;""")        
        print("||  В какой день будет урок?")
        day = input("  ")
        print("||  У какого класса?")
        grade = input("  ")
        print("||  Какой номер урока?")
        n_lesson = int(input("  "))
        request = """
            SELECT room FROM teachers
            WHERE (day = ? AND grade = ?) AND n_lesson = ?;
        """
        cur.execute(request, (day, grade, n_lesson))
        #print (output)
        output = cur.fetchall()
        if output == []:
            print("||  Информация не найдена")
        else:
            for row in output:
                print ("\t", '%-10s' % (row[0]))
        os.system('pause')
    except sq.Error as e:
        if db: db.rollback()
        print("||  Ошибка выполнения запроса ", e)
        os.system('pause')
    finally:
        if db: db.close()
def key6():
    try:
        db = sq.connect("school.db")
        cur = db.cursor()
        cur.execute("""Begin;""")        
        print ("||  Введите ФИО преподавателя")
        name_1 = input("  ")
        name_2 = input("  ")
        name_3 = input("  ")
        print ("||  Введите название предмета")
        subject = input("  ")
        request = """
            SELECT DISTINCT grade
            FROM teachers
            JOIN id_teachers ON id_teachers.id_teacher = teachers.id_teacher
            WHERE (name_1 = ? AND name_2 = ? AND name_3 = ?) AND subject = ?
            ORDER BY grade
        """
        cur.execute(request, (name_1, name_2, name_3, subject))
        #print (output)
        output = cur.fetchall()
        if output == []:
            print("||  Информация не найдена")
        else:
            print("|| ", name_2, name_3, "преподает в:")
            for row in output:
                print ("\t", '%-4s' % (row[0]))
        os.system('pause')
    except sq.Error as e:
        if db: db.rollback()
        print("||  Ошибка выполнения запроса ", e)
        os.system('pause')
    finally:
        if db: db.close()
def key7():
    try:
        print("||  Формирование справки о количестве учеников в классе\n")
        db = sq.connect("school.db")
        cur = db.cursor()
        cur.execute("""Begin;""")
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        head = doc.add_heading("Cправка о количестве учеников", level=0)
        head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        print("||  Введите класс")
        grade = input("  ")
        doc_text = doc.add_paragraph('В ')
        doc_text.add_run(grade).bold = True
        doc_text.add_run(' учатся ')
        request = """
            SELECT count(id_student) as "Количество учеников"
            FROM id_students
            WHERE grade = ?
        """
        cur.execute(request, (grade,))
        #print (output)
        output = cur.fetchall()
        if output == []:
            print("||  Такого класса нет")
        else:
            for row in output:
                print ("\t", '%-5s' % (row[0]))
                print("||  Количество учеников в", grade) 
                doc_text.add_run(str(row[0]))
                doc_text.add_run(' учеников.')
                doc.save("StudentsCount.docx")
        os.system('pause')
    except sq.Error as e:
        if db: db.rollback()
        print("||  Ошибка выполнения запроса ", e)
        os.system('pause')
    finally:
        if db: db.close()
def key8():
    try:
        print("||  Формирование отчета о работе школы\n")
        db = sq.connect("school.db")
        cur = db.cursor()
        cur.execute("""Begin;""")
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        head = doc.add_heading("Отчет о работе школы", level=0)
        head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        #1 - количество учителей по предметам
        doc_text = doc.add_paragraph('Количество учителей по предметам:')
        doc.add_paragraph()
        request = """
            SELECT subject as "Предмет", count(DISTINCT id_teacher) as "Количество учителей"
            FROM teachers group by subject;
        """
        cur.execute(request)
        #print (output)
        output = cur.fetchall()
        table = doc.add_table(rows=1, cols=2)
        cells = table.rows[0].cells
        cells[0].text = 'Предмет'
        cells[1].text = 'Количество учителей'
        for subj, n_teach in output:
            #print ("\t", '%-15s %-5s' % (subj,n_teach))
            cells = table.add_row().cells
            cells[0].text = subj
            cells[1].text = str(n_teach)
        doc.add_paragraph()
        #2 - количество кабинетов
        doc_text = doc.add_paragraph('Количество кабинетов в школе: ')
        request = """
            SELECT count(DISTINCT room) FROM teachers
            WHERE not (room = "");
        """
        cur.execute(request)
        #print (output)
        output = cur.fetchall()
        for row in output:
            #print ("\t", '%-5s' % (row[0]))
            doc_text.add_run(str(row[0]))
        doc_text.add_run('.')
        doc.add_page_break()
        #3 - Количество учителей в каждом классе
        doc_text = doc.add_paragraph('Количество учителей на класс: ')
        doc.add_paragraph()
        request = """
            select grade as "Класс", count(DISTINCT id_teacher) as "Кол-во учителей"
            from teachers group by grade;
        """
        cur.execute(request)
        #print (output)
        output = cur.fetchall()
        table = doc.add_table(rows=1, cols=2)
        cells = table.rows[0].cells
        cells[0].text = 'Класс'
        cells[1].text = 'Количество учителей'
        for grade, n_teach in output:
            #print ("\t", '%-5s %-5s' % (grade, n_teach))
            cells = table.add_row().cells
            cells[0].text = grade
            cells[1].text = str(n_teach)
        #4 - количество двоечников
        doc_text = doc.add_paragraph('Количество двоечников в школе: ')
        request = """
            SELECT count(DISTINCT id_student) as "Кол-во учеников" from students
        """
        cur.execute(request)
        #print (output)
        output = cur.fetchall()
        for row in output:
            #print ("\t", '%-5s' % (row[0]))
            doc_text.add_run(str(row[0]))
        doc_text.add_run('.')
        #5 - количество хорошистов
        doc_text = doc.add_paragraph('Количество хорошистов в школе: ')
        request = """
            SELECT count(DISTINCT id_student) as "Кол-во учеников" from students
        """
        cur.execute(request)
        #print (output)
        output = cur.fetchall()
        for row in output:
            #print ("\t", '%-5s' % (row[0]))
            doc_text.add_run(str(row[0]))
        doc_text.add_run('.')
        #6 - количество отличников
        doc_text = doc.add_paragraph('Количество отличников в школе: ')
        request = """
            SELECT count(DISTINCT id_student) as "Кол-во учеников" from students
        """
        cur.execute(request)
        #print (output)
        output = cur.fetchall()
        for row in output:
            #print ("\t", '%-5s' % (row[0]))
            doc_text.add_run(str(row[0]))
        doc_text.add_run('.')
        doc.save("School.docx")
        os.system('pause')
    except sq.Error as e:
        if db: db.rollback()
        print("||  Ошибка выполнения запроса ", e)
        os.system('pause')
    finally:
        if db: db.close()

def key9():
    print("Создать резервную копию БД или восстановить БД из файла?\n")
    print("1 - Создать")
    print("2 - Восстановить")
    print("0 - Вернуться назад\n")
    key = input()
    while key != '0':
        if key == '1':
            os.system('cls')
            sq_dump_w()
            print("Резервная копия создана\n")
            os.system('pause')
            break
        if key == '2':
            os.system('cls')
            sq_dump_r()
            print("Резервная копия восстановлена\n")
            os.system('pause')
            break

def main():
    sq_init()
    choose_user()
    user = input()
    choose_main(user)
    key = input()
    while key != '0':
        if key == '1':
            os.system('cls')
            key1()
            choose_main(user)
            key = input()
        if key == '2' and user == '2':
            os.system('cls')
            key2()
            choose_main(user)
            key = input()
        if key == '3':
            os.system('cls')
            key3()
            choose_main(user)
            key = input()
        if key == '4' and user == '2':
            os.system('cls')
            key4()
            choose_main(user)
            key = input()
        if key == '5':
            os.system('cls')
            key5()
            choose_main(user)
            key = input()
        if key == '6' and user == '2':
            os.system('cls')
            key6()
            choose_main(user)
            key = input()
        if key == '7' and user == '2':
            os.system('cls')
            key7()
            choose_main(user)
            key = input()
        if key == '8' and user == '2':
            os.system('cls')
            key8()
            choose_main(user)
            key = input()
        if key == '9' and user == '2':
            os.system('cls')
            key9()
            choose_main(user)
            key = input()
    return 0
    
main()